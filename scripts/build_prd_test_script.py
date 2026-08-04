from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path("artifacts/Kich_ban_kiem_thu_PRD_WebShadowing.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
NAVY = "1F4D78"
LIGHT = "E8EEF5"
PALE = "F5F8FC"
GRAY = "666666"
RED = "C62828"
GREEN = "2E7D32"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    total = int(sum(widths) * 1440)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(width * 1440)))
        grid.append(gridCol)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            w = int(widths[idx] * 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar, instr, fldChar2])

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Title", 26, NAVY, 0, 10),
    ("Subtitle", 13, GRAY, 0, 18),
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = name != "Subtitle"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for sty in ["List Bullet", "List Number"]:
    s = styles[sty]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.paragraph_format.left_indent = Inches(0.375)
    s.paragraph_format.first_line_indent = Inches(-0.188)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.25

if "TC Label" not in styles:
    st = styles.add_style("TC Label", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "Calibri"
    st.font.size = Pt(9)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(GRAY)
    st.paragraph_format.space_after = Pt(2)

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.add_run("WEB SHADOWING  |  KỊCH BẢN KIỂM THỬ PRD"), 8.5, True, GRAY)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(fp.add_run("Tài liệu kiểm thử chức năng & phi chức năng  •  Trang "), 8.5, False, GRAY)
add_field(fp, "PAGE")
set_font(fp.add_run(" / "), 8.5, False, GRAY)
add_field(fp, "NUMPAGES")

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("KỊCH BẢN KIỂM THỬ CHI TIẾT\nWEB SHADOWING")
sub = doc.add_paragraph(style="Subtitle")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Theo luồng người dùng PRD • Functional • Responsive • Non-functional")

meta = doc.add_table(rows=5, cols=2)
meta.style = "Table Grid"
set_table_widths(meta, [1.875, 4.625])
for i, (a, b) in enumerate([
    ("Phiên bản", "1.0"),
    ("Ngày lập", "30/07/2026"),
    ("Ngôn ngữ", "Tiếng Việt"),
    ("Phạm vi", "Web responsive: Desktop / Tablet / Mobile; tài khoản Khách / Free / VIP"),
    ("Nguồn yêu cầu", "PRD do người yêu cầu cung cấp trong phiên làm việc"),
]):
    meta.cell(i, 0).text = a
    meta.cell(i, 1).text = b
    set_cell_shading(meta.cell(i, 0), LIGHT)
    for p in meta.cell(i, 0).paragraphs:
        for r in p.runs: set_font(r, 10, True, NAVY)
    for p in meta.cell(i, 1).paragraphs:
        for r in p.runs: set_font(r, 10)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Mục tiêu: kiểm tra đầy đủ hành trình từ xác thực, chọn Mode, học bài, chấm điểm, quản lý tiến trình đến tài khoản/VIP; không thực thi kiểm thử."), 10.5, False, GRAY, True)
doc.add_page_break()

doc.add_heading("1. Cách sử dụng tài liệu", level=1)
for text in [
    "Mỗi test case độc lập, trừ khi phần Tiền điều kiện nêu rõ phụ thuộc vào test case khác.",
    "Ghi nhận Pass khi tất cả kết quả mong đợi đều đúng; chỉ một sai lệch cũng ghi Fail và tạo defect kèm ảnh/video, log, tài khoản, thiết bị, trình duyệt và thời điểm.",
    "Các giá trị chưa được PRD chốt được đánh dấu [CẦN PO XÁC NHẬN]. Tester thay bằng cấu hình thực tế trước khi chạy.",
    "Với thao tác làm thay đổi EXP/Tim/Streak/Mode/VIP, phải đối chiếu đồng thời Top Nav, trang Tiến trình & Thẻ nhớ, dữ liệu sau refresh và dữ liệu sau đăng nhập lại.",
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("2. Giả định và điểm cần Product xác nhận", level=1)
assumptions = [
    ("A-01", "Trang mặc định", "Sau đăng nhập/đăng ký thành công mở Trang Khóa Học; câu “dẫn vào Trang chủ” được hiểu là Trang Khóa Học."),
    ("A-02", "Breakpoint", "Desktop ≥ 1024 px; Tablet 768–1023 px; Mobile ≤ 767 px. [CẦN PO XÁC NHẬN]"),
    ("A-03", "Đổi Mode Free", "Giới hạn 1 lần/tháng hoặc mất EXP; tài liệu kiểm cả hai phương án cấu hình. [CẦN PO CHỐT]"),
    ("A-04", "EXP đổi Tim", "Mặc định 100 EXP = 1 Tim; tỷ lệ phải cấu hình được để kiểm thử chống lạm phát."),
    ("A-05", "Số lần phát âm sai", "“> 3 lần liên tiếp” nghĩa là tự lưu ở lần sai thứ 4; một lần đạt chuẩn sẽ reset chuỗi sai."),
    ("A-06", "Nghe chép", "Bỏ qua hoa/thường và dấu câu; khoảng trắng thừa được chuẩn hóa. Ngưỡng đạt dùng 50/70/90% đã cấu hình."),
    ("A-07", "Tim Free", "Khi Tim = 0, không cho bắt đầu/chấm câu có thể trừ Tim và hiển thị hướng dẫn đổi EXP/mua VIP. [CẦN PO XÁC NHẬN]"),
    ("A-08", "Streak", "Tăng tối đa 1 lần/ngày khi hoàn thành ít nhất một hoạt động hợp lệ. Múi giờ Asia/Saigon. [CẦN PO XÁC NHẬN]"),
    ("A-09", "AI không lưu", "Bài AI chưa lưu bị mất khi rời trang/đăng xuất; nếu auto-save bật thì lưu ngay khi tạo thành công."),
    ("A-10", "Xóa", "Xóa bài AI/câu yêu thích/thẻ nhớ có xác nhận; thao tác thành công phải bền vững sau refresh."),
]
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
set_table_widths(t, [0.7, 1.35, 4.45])
for i, h in enumerate(["Mã", "Chủ đề", "Giả định / câu hỏi"]):
    t.cell(0, i).text = h
    set_cell_shading(t.cell(0, i), LIGHT)
    for r in t.cell(0, i).paragraphs[0].runs: set_font(r, 9.5, True, NAVY)
set_repeat_table_header(t.rows[0])
for row in assumptions:
    cells = t.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val
        for r in cells[i].paragraphs[0].runs: set_font(r, 9.5)
    set_table_widths(t, [0.7, 1.35, 4.45])

doc.add_heading("3. Ma trận môi trường và dữ liệu", level=1)
for text in [
    "Desktop: 1920×1080, 1440×900, 1366×768, 1024×768; zoom 80%, 100%, 125%, 200%.",
    "Tablet: 1024×768 landscape, 768×1024 portrait.",
    "Mobile: 393×852, 390×844, 375×812, 360×800, 320×568; xoay portrait/landscape.",
    "Trình duyệt: Chrome, Edge, Firefox bản ổn định mới nhất; Safari iOS/macOS nếu thuộc phạm vi phát hành.",
    "Mạng: online ổn định, Slow 3G, mất mạng giữa request, timeout API, retry, offline.",
    "Tài khoản: Khách; Free mới; Free có lịch sử; Free 0/1/nhiều Tim; Free 99/100/101 EXP; VIP đang hoạt động; VIP hết hạn/hủy gia hạn.",
    "Quyền thiết bị: Mic Allow / Block / Ask every time / bị hệ điều hành chặn; loa muted; không có thiết bị input.",
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("4. Quy tắc kiểm tra xuyên suốt", level=1)
rules = [
    ("EXP", "Mỗi câu đúng chỉ +20 một lần; retry/refresh/back-forward/double click không cộng lặp."),
    ("Tim Free", "Mỗi lần không đạt/sai ở Tab 1, 3, 4 chỉ trừ 1; không trừ ở thao tác không gửi đáp án, lỗi API hoặc lỗi mạng."),
    ("Tim VIP", "Hiển thị vô hạn; không bị trừ dù câu sai."),
    ("Ngưỡng", "Điểm = ngưỡng cho qua; điểm < ngưỡng không cho qua; điểm > ngưỡng cho qua."),
    ("Mode", "Casual ưu tiên fluency/ngữ điệu và chấp nhận slang; Academic nghiêm phoneme/trọng âm/ngữ pháp; Professional ưu tiên clarity/sắc thái lịch sự."),
    ("Đồng bộ", "Top Nav, trang Tiến trình, server và phiên đăng nhập mới phải cùng số liệu."),
    ("Idempotency", "Double click, retry mạng và callback lặp không tạo giao dịch/EXP/Tim/bài lưu trùng."),
]
rt = doc.add_table(rows=1, cols=2)
rt.style = "Table Grid"
set_table_widths(rt, [1.181, 5.319])
for i,h in enumerate(["Đối tượng","Bất biến cần giữ"]):
    rt.cell(0,i).text=h; set_cell_shading(rt.cell(0,i), LIGHT)
    for r in rt.cell(0,i).paragraphs[0].runs: set_font(r,10,True,NAVY)
for a,b in rules:
    c=rt.add_row().cells; c[0].text=a; c[1].text=b
    for r in c[0].paragraphs[0].runs: set_font(r,9.5,True)
    for r in c[1].paragraphs[0].runs: set_font(r,9.5)
    set_table_widths(rt,[1.181,5.319])

cases = []
def tc(area, title, steps, expected, pre="Ứng dụng hoạt động; dữ liệu theo mô tả.", data="Không có", priority="P1", kind="Chức năng"):
    cases.append(dict(area=area,title=title,steps=steps,expected=expected,pre=pre,data=data,priority=priority,kind=kind))

# 01 Authentication & onboarding
tc("Xác thực & Onboarding","Đăng nhập hợp lệ và điều hướng mặc định",
   ["Mở trang đăng nhập.","Nhập email và mật khẩu hợp lệ.","Bấm nút “Đăng nhập” một lần."],
   ["Hiển thị trạng thái đang xử lý; nút chống bấm lặp.","Đăng nhập thành công và mở Trang Khóa Học.","Top Nav/Sub-nav phản ánh đúng tài khoản; không quay lại trang đăng nhập bằng Back nếu phiên còn hợp lệ."],
   data="Free hợp lệ")
tc("Xác thực & Onboarding","Đăng nhập sai mật khẩu",
   ["Nhập email hợp lệ và mật khẩu sai.","Bấm “Đăng nhập”."],
   ["Không tạo phiên đăng nhập; ở lại form.","Thông báo lỗi rõ ràng, không tiết lộ tài khoản có tồn tại; mật khẩu không hiện trong URL/log/UI.","Cho phép sửa và gửi lại."],data="Email hợp lệ; mật khẩu sai")
tc("Xác thực & Onboarding","Validation form đăng nhập",
   ["Để trống cả hai trường và bấm “Đăng nhập”.","Nhập email sai định dạng.","Nhập chuỗi chỉ có khoảng trắng; thử nhấn Enter."],
   ["Lỗi hiển thị sát đúng trường; focus tới lỗi đầu tiên.","Không gửi request khi dữ liệu client không hợp lệ.","Enter có hành vi giống nút Đăng nhập; không reload trắng trang."],priority="P2")
tc("Xác thực & Onboarding","Đăng nhập bị timeout/mất mạng",
   ["Nhập thông tin hợp lệ.","Ngắt mạng hoặc ép request timeout.","Bấm “Đăng nhập”; sau đó khôi phục mạng và bấm thử lại."],
   ["Hiển thị lỗi có thể thử lại; không treo loading.","Không tạo phiên nửa vời/đăng nhập kép.","Thử lại thành công điều hướng đúng."],kind="Khả năng phục hồi")
tc("Xác thực & Onboarding","Mở luồng đăng ký",
   ["Tại đăng nhập, bấm “Đăng ký”.","Quan sát form và điều hướng quay lại."],
   ["Mở đúng form đăng ký; tiêu đề, trường bắt buộc, chính sách và nút rõ ràng.","Bấm “Đăng nhập”/Back quay lại không làm vỡ trạng thái."],priority="P2")
tc("Xác thực & Onboarding","Đăng ký hợp lệ – Mode Giao tiếp – ngưỡng 50%",
   ["Điền thông tin tài khoản hợp lệ.","Bấm “Tiếp tục”.","Chọn card “Giao tiếp”.","Chọn “50%”.","Bấm nút hoàn tất đăng ký."],
   ["Chỉ một Mode và một ngưỡng được chọn.","Tài khoản tạo đúng; cấu hình Giao tiếp/50% được lưu.","Hiển thị quảng cáo VIP theo thiết kế rồi có hành động tiếp tục/bỏ qua; cuối luồng vào Trang Khóa Học."],data="Email mới")
tc("Xác thực & Onboarding","Đăng ký hợp lệ – Mode Học thuật – ngưỡng 70%",
   ["Đăng ký tài khoản mới.","Chọn “Học thuật”.","Chọn “70%”.","Hoàn tất."],
   ["Cấu hình Học thuật/70% được lưu và hiển thị đúng trong Tài khoản.","Nội dung khóa học mặc định lọc theo Học thuật."],data="Email mới")
tc("Xác thực & Onboarding","Đăng ký hợp lệ – Mode Công việc – ngưỡng 90%",
   ["Đăng ký tài khoản mới.","Chọn “Công việc”.","Chọn “90%”.","Hoàn tất."],
   ["Cấu hình Công việc/90% được lưu.","Nội dung mặc định đúng Mode; tiêu chuẩn 90% áp dụng cho Tab 1 và Tab 3."],data="Email mới")
tc("Xác thực & Onboarding","Không chọn Mode hoặc ngưỡng",
   ["Điền tài khoản hợp lệ.","Không chọn Mode và bấm tiếp tục.","Chọn Mode nhưng không chọn ngưỡng rồi bấm hoàn tất."],
   ["Không hoàn tất đăng ký.","Thông báo bắt buộc đúng nhóm điều khiển và focus hợp lý.","Không tạo tài khoản/cấu hình dở dang ngoài chủ đích."],priority="P1")
tc("Xác thực & Onboarding","Đổi lựa chọn trước khi hoàn tất",
   ["Chọn Giao tiếp rồi chọn Học thuật.","Chọn 50% rồi chọn 90%.","Quay lại bước trước và tiến tới lại."],
   ["Lựa chọn sau cùng là duy nhất và được giữ đúng.","UI selected/aria-checked đồng nhất.","Sau đăng ký lưu Học thuật/90%, không lưu giá trị cũ."],priority="P2")
tc("Xác thực & Onboarding","Quảng cáo VIP – đóng/bỏ qua",
   ["Hoàn tất cấu hình đăng ký.","Tại quảng cáo VIP, bấm “Bỏ qua” hoặc nút đóng.","Refresh sau khi vào app."],
   ["Đóng được mà không mua; vào Khóa Học với Free.","Không tự kích hoạt VIP/không mở cổng thanh toán.","Không lặp quảng cáo bất thường theo cùng phiên nếu thiết kế không yêu cầu."],priority="P2")
tc("Xác thực & Onboarding","Quảng cáo VIP – chọn mua",
   ["Hoàn tất đăng ký.","Bấm CTA “Mua VIP/Xem gói”.","Dùng Back/Đóng tại bước thanh toán."],
   ["Mở đúng luồng/gói/giá; chưa thanh toán thì vẫn Free.","Quay lại an toàn, không mất tài khoản/cấu hình onboarding."],priority="P1")

# Navigation
nav_items=[("Logo","Bấm Logo từ trang con","Mở Trang Khóa Học; không mất phiên."),
("Khóa học","Bấm mục Khóa học","Mục active đúng và nội dung khóa học hiển thị."),
("Tiến trình & Thẻ nhớ","Bấm mục Tiến trình & Thẻ nhớ","Mở đúng trang và active state đúng."),
("Tài khoản","Bấm mục Tài khoản","Mở Settings; dữ liệu tài khoản đúng.")]
for name,step,exp in nav_items:
    tc("Navigation",f"Điều hướng bằng {name}",[step,"Bấm Back rồi Forward của trình duyệt."],[exp,"URL/lịch sử trình duyệt đúng; focus/active state cập nhật; không reload mất dữ liệu đã lưu."],priority="P2")
tc("Navigation","Top Nav Desktop đúng thứ tự và dữ liệu",
   ["Đặt viewport 1440×900 và đăng nhập.","Đọc từ trái sang phải toàn bộ Top Nav.","Thay đổi EXP/Tim ở hoạt động hợp lệ và quan sát."],
   ["Thứ tự: Logo → Streak → Tim → EXP → loại gói → công cụ Dark/Light → Đăng xuất/Đăng nhập.","Không chồng lấn/cắt chữ; số liệu cập nhật đúng và có nhãn/tooltip dễ hiểu.","Khách thấy Đăng nhập; người dùng thấy Đăng xuất."],kind="Responsive")
tc("Navigation","Mobile Sub-nav và dời công cụ",
   ["Đặt viewport 390×844.","Quan sát thanh điều hướng.","Mở Tài khoản."],
   ["Sub-nav chỉ gồm Khóa học | Tiến trình & Thẻ nhớ | Tài khoản.","Dark/Light và Đăng xuất không chiếm Top Nav; xuất hiện trong Tài khoản.","Không có scroll ngang; vùng bấm tối thiểu 44×44 CSS px."],kind="Responsive")
tc("Navigation","Chuyển Dark/Light trên Desktop",
   ["Bấm công tắc Dark/Light ở Top Nav.","Đi qua ba trang; refresh.","Đăng xuất rồi đăng nhập lại."],
   ["Theme đổi ngay toàn bộ nền/chữ/icon/form/modal; tương phản đọc được.","Không chớp theme sai kéo dài; lựa chọn được lưu theo thiết kế.","Không thay đổi dữ liệu học."],priority="P1")
tc("Navigation","Đăng xuất Desktop",
   ["Đăng nhập và mở trang bất kỳ.","Bấm “Đăng xuất” trên Top Nav.","Xác nhận nếu có; dùng Back."],
   ["Xóa phiên/token an toàn; chuyển về trạng thái khách/đăng nhập.","Trang bảo vệ không còn truy cập qua Back/URL trực tiếp.","Nút đổi thành Đăng nhập; không lộ dữ liệu người dùng cũ."],priority="P0")
tc("Navigation","Phiên hết hạn khi đang thao tác",
   ["Đăng nhập rồi ép token hết hạn.","Mở trang bảo vệ hoặc gửi đáp án.","Đăng nhập lại."],
   ["Thông báo phiên hết hạn, chuyển đăng nhập; không vòng lặp redirect.","Không trừ Tim/cộng EXP với request thất bại do hết phiên.","Sau đăng nhập lại, dữ liệu server đúng."],priority="P0",kind="Bảo mật/khả năng phục hồi")

# Course page
for typ in ["Giáo trình","Video Bank"]:
    for mode in ["Giao tiếp","Học thuật","Công việc"]:
        tc("Trang Khóa Học",f"{typ} – lọc Mode {mode}",
           ["Mở Trang Khóa Học.",f"Chọn tab/nhóm “{typ}”.",f"Chọn bộ lọc “{mode}”.","Mở card đầu tiên."],
           [f"Chỉ hiển thị {typ} thuộc {mode}; trạng thái bộ lọc rõ ràng.","Card hiển thị tối thiểu tên, phân loại, Unit/thông tin tiến độ phù hợp.","Mở đúng bài; Back giữ bộ lọc/vị trí cuộn hợp lý."],priority="P1")
tc("Trang Khóa Học","Card Unit – trạng thái và thao tác",
   ["Quan sát card chưa học, đang học, đã hoàn thành.","Bấm toàn bộ vùng card và CTA.","Dùng bàn phím Tab + Enter."],
   ["Trạng thái/tiến độ phân biệt bằng chữ và không chỉ bằng màu.","Chỉ mở một bài; không điều hướng kép.","Có focus visible; Enter kích hoạt giống click."],priority="P2",kind="UI/Accessibility")
tc("Trang Khóa Học","Danh sách rỗng và lỗi tải học liệu",
   ["Chọn bộ lọc không có dữ liệu.","Giả lập API 500 rồi bấm tải lại."],
   ["Empty state đúng ngữ cảnh, có hành động đổi bộ lọc/tạo AI.","Lỗi không hiển thị trang trắng; nút Thử lại hoạt động; không nhân đôi card."],kind="Khả năng phục hồi")
tc("Bài học AI","Tạo bài AI Mode Giao tiếp",
   ["Chọn tạo bài AI.","Nhập prompt về sở thích/đời sống hằng ngày.","Chọn vai bạn bản xứ/bạn game/hướng dẫn viên.","Bấm “Tạo bài học”."],
   ["Có validation và loading; chống gửi lặp.","AI sinh phụ đề và audio tự nhiên, không video; nội dung bám Mode/vai/prompt.","Mặc định chưa lưu Database; có nút “Lưu bài học” ngay sau tạo."],data="Prompt hợp lệ ≤ giới hạn")
tc("Bài học AI","Tạo bài AI Mode Học thuật",
   ["Chọn Mode Học thuật.","Nhập prompt khoa học/xã hội/thuyết trình.","Chọn vai giáo viên/điều phối/giám khảo IELTS.","Bấm tạo."],
   ["Nội dung đúng vai và phong cách học thuật; có phụ đề + audio.","Không lẫn tình huống Casual/Professional; không tự lưu khi auto-save tắt."])
tc("Bài học AI","Tạo bài AI Mode Công việc",
   ["Chọn Mode Công việc.","Nhập prompt báo cáo tiến độ/thương lượng/khiếu nại.","Chọn vai nhà tuyển dụng/khách hàng/quản lý dự án.","Bấm tạo."],
   ["Nội dung đúng bối cảnh chuyên nghiệp và prompt; phụ đề/audio dùng được.","Không tự thêm video; CTA lưu hiển thị."])
tc("Bài học AI","Validation prompt AI",
   ["Để prompt trống/chỉ khoảng trắng và bấm tạo.","Nhập prompt vượt giới hạn.","Nhập HTML/script và ký tự Unicode/emoji."],
   ["Trống/quá dài bị chặn với thông báo và bộ đếm rõ ràng.","Nội dung được encode/sanitize; không thực thi script.","Không gửi request không hợp lệ; dữ liệu nhập không làm vỡ layout."],priority="P0",kind="Validation/Bảo mật")
tc("Bài học AI","Lỗi/timeout/kiểm duyệt khi tạo AI",
   ["Gửi prompt rồi ép timeout/API 500.","Thử lại.","Gửi nội dung bị chính sách từ chối."],
   ["Loading kết thúc; thông báo phù hợp, không mất prompt.","Retry không tạo bài trùng/lưu trùng.","Từ chối an toàn, không hiển thị nội dung độc hại hay stack trace."],kind="Khả năng phục hồi/Bảo mật")
tc("Bài học AI","Lưu bài ngay sau khi tạo",
   ["Tạo bài khi auto-save tắt.","Bấm “Lưu bài học” hai lần nhanh.","Refresh và mở danh sách bài AI."],
   ["Lưu đúng một bản; nút chuyển trạng thái Đã lưu/disabled.","Bài tồn tại sau refresh, đúng title/prompt/mode/audio/phụ đề.","Không trừ EXP/Tim ngoài PRD."],priority="P1")
tc("Bài học AI","Lưu bài khi đang học",
   ["Tạo bài nhưng chưa lưu; bắt đầu học.","Bấm “Lưu bài học” trong màn học.","Rời trang rồi mở lại."],
   ["Lưu đúng bài đang học và không reset tiến độ hiện tại.","Mở lại được; không tạo bản trùng."],priority="P1")
tc("Bài học AI","Auto-save bật/tắt",
   ["Trong Tài khoản bật auto-save và lưu.","Tạo bài AI A; kiểm danh sách.","Tắt auto-save; tạo bài B; kiểm danh sách."],
   ["A tự lưu đúng một bản sau tạo thành công.","B không tự lưu; chỉ lưu khi bấm nút.","Thiết lập bền vững sau refresh/đăng nhập lại."],priority="P1")
tc("Bài học AI","Xóa bài AI – xác nhận/hủy",
   ["Tại bài đã lưu bấm “Xóa”.","Trong popup bấm “Hủy”.","Bấm Xóa lại và xác nhận.","Refresh."],
   ["Hủy giữ nguyên bài; focus trở lại nút gọi modal.","Xác nhận xóa đúng một bài, thông báo thành công.","Bài không còn sau refresh; URL cũ xử lý 404/empty hợp lý."],priority="P1")
tc("Bài học AI","Rời bài AI chưa lưu",
   ["Tạo bài khi auto-save tắt.","Bấm Back/chuyển trang/đóng tab.","Chọn ở lại rồi lặp lại và chọn rời."],
   ["Có cảnh báo mất nội dung nếu thiết kế yêu cầu; “Ở lại” giữ bài.","“Rời” không lưu vào Database; không để bản ghi rác.","Không cảnh báo khi bài đã lưu."],priority="P2")

# Lesson tabs
tc("Bài giảng – Chung","Hiển thị và chuyển 4 Tabs",
   ["Mở một bài giảng.","Lần lượt bấm Shadowing Studio, Đối thoại với AI, Nghe chép chính tả, Ghép IPA.","Dùng bàn phím Left/Right nếu tablist hỗ trợ."],
   ["Đúng 4 tab, một tab active; nội dung không trộn.","Chuyển tab không cộng EXP/trừ Tim; trạng thái bài được giữ theo thiết kế.","Tab có role/name/state và focus visible."],priority="P1",kind="Chức năng/Accessibility")
tc("Tab 1 – Shadowing","Phát/Tạm dừng/Tua/âm lượng Audio-Video",
   ["Bấm Play.","Bấm Pause, tua timeline, đổi âm lượng/mute, phát lại.","Chuyển câu rồi quay lại."],
   ["Media phát đúng bài/câu; icon và thời gian cập nhật.","Không tự phát trái chính sách trình duyệt; điều khiển bàn phím được.","Audio có ngữ điệu tự nhiên, không méo/đứt; transcript đồng bộ hợp lý."],priority="P1")
tc("Tab 1 – Shadowing","Transcript, phiên âm, dịch nghĩa và bấm từ",
   ["Quan sát từng từ trong transcript.","Bấm một từ; bấm từ khác; bấm ngoài popup.","Kiểm tra từ dài, dấu nháy, dạng rút gọn."],
   ["IPA của mỗi từ nằm ngay dưới đúng từ; không lệch khi wrap.","Bấm từ mở nghĩa đúng, từ khác cập nhật popup, bấm ngoài đóng.","Dịch nghĩa câu hiển thị đúng; từ là nút/focus được, không chỉ hover."],priority="P1")
tc("Tab 1 – Shadowing","Cấp quyền mic lần đầu",
   ["Bấm nút Mic/Ghi âm.","Ở prompt trình duyệt chọn Allow.","Nói và bấm Dừng/Gửi chấm."],
   ["Giải thích mục đích trước/đúng lúc xin quyền; mic active có chỉ báo rõ.","Thu đúng audio; chỉ gửi khi người dùng yêu cầu.","Loading chấm điểm và kết quả hiển thị; mic dừng sau hoàn tất/rời trang."],priority="P0",kind="Quyền riêng tư")
tc("Tab 1 – Shadowing","Từ chối quyền mic",
   ["Bấm Mic và chọn Block.","Bấm thử lại.","Mở hướng dẫn cấp quyền."],
   ["Không crash/thu âm giả; thông báo nguyên nhân và cách bật lại.","Không trừ Tim vì chưa có bài chấm.","Nút/luồng thay thế rõ ràng; không spam prompt quyền."],priority="P0")
tc("Tab 1 – Shadowing","Không có mic hoặc mic đang bận",
   ["Chọn thiết bị không có input hoặc chiếm mic bằng app khác.","Bấm Ghi âm."],
   ["Lỗi thân thiện, không treo waveform/loading.","Không trừ Tim/cộng EXP; không gửi file rỗng.","Cho phép thử lại sau khi thiết bị sẵn sàng."],kind="Khả năng phục hồi")
tc("Tab 1 – Shadowing","Điểm đúng tại biên 50%",
   ["Đặt mục tiêu 50%.","Gửi mẫu được mock chấm 49%.","Gửi mẫu 50%; sau đó 51%."],
   ["49%: không Next, trừ đúng 1 Tim Free.","50% và 51%: cho Next; mỗi lần hoàn thành câu hợp lệ +20 EXP đúng một lần.","Hiển thị điểm/feedback khớp response API."],data="Mock score 49/50/51; Free")
tc("Tab 1 – Shadowing","Điểm đúng tại biên 70%",
   ["Đặt mục tiêu 70%.","Chấm 69%, 70%, 71% ở các câu độc lập."],
   ["69% fail/trừ 1 Tim; 70% và 71% pass/cho Next/+20 EXP mỗi câu.","Ngưỡng hiển thị và logic cùng là 70%."],data="Mock score 69/70/71; Free")
tc("Tab 1 – Shadowing","Điểm đúng tại biên 90%",
   ["Đặt mục tiêu 90%.","Chấm 89%, 90%, 91% ở các câu độc lập."],
   ["89% fail/trừ 1 Tim; 90% và 91% pass/cho Next/+20 EXP.","Không làm tròn 89.x thành pass trái quy tắc nếu API trả số thập phân."],data="Mock score 89/90/91; Free")
tc("Tab 1 – Shadowing","Sai một từ 4 lần liên tiếp tự vào Sổ tay",
   ["Chọn một câu có từ target.","Gửi 4 lần liên tiếp đều đánh dấu target sai.","Mở Tiến trình & Thẻ nhớ."],
   ["Lần 1–3 chưa tự thêm; lần 4 tự thêm đúng một flashcard.","Mỗi lần fail trừ 1 Tim Free; không thêm trùng ở lần 5.","Flashcard có từ/IPA/nghĩa/nguồn phù hợp."],data="Target pronunciation fail x4")
tc("Tab 1 – Shadowing","Reset chuỗi sai khi phát âm đạt",
   ["Phát âm target sai 3 lần.","Lần 4 đạt chuẩn.","Tiếp tục sai 1 lần; kiểm Sổ tay."],
   ["Lần đạt reset chuỗi sai liên tiếp.","Một lần sai mới không tự thêm từ; chỉ tự thêm khi đạt lại 4 sai liên tiếp."],data="Fail x3 → Pass → Fail")
tc("Tab 1 – Shadowing","Lưu câu thoại yêu thích",
   ["Bấm icon/nút “Lưu câu” ở câu hiện tại.","Bấm lại nhanh hoặc bấm lại sau khi đã lưu.","Mở danh sách yêu thích."],
   ["Có feedback trạng thái đã lưu; đúng một bản ghi.","Câu, dịch, nguồn bài và thứ tự hiển thị đúng.","Không ảnh hưởng EXP/Tim; trạng thái giữ sau refresh."],priority="P1")
tc("Tab 1 – Shadowing","API đánh giá phát âm lỗi/timeout",
   ["Thu âm hợp lệ.","Ép API chấm 500/timeout/malformed response.","Bấm Thử lại."],
   ["Không dùng Web STT thay thế để tạo điểm giả.","Không trừ Tim/cộng EXP khi không có kết quả hợp lệ; audio/xử lý quyền riêng tư đúng.","Retry tạo đúng một kết quả, không giao dịch lặp."],priority="P0",kind="Tích hợp/Khả năng phục hồi")
tc("Tab 1 – Shadowing","Double click Gửi chấm/Next",
   ["Thu âm xong.","Double click “Gửi chấm”.","Sau pass, double click “Next”."],
   ["Chỉ một request chấm; nút disabled/loading.","Chỉ +20 EXP một lần; chỉ chuyển một câu.","Không bỏ qua câu kế tiếp/không nhân lịch sử."],priority="P0")

tc("Tab 2 – AI Voice Chat","Free truy cập tính năng VIP",
   ["Đăng nhập Free và mở bài.","Bấm tab “Đối thoại với AI”."],
   ["Không cho dùng chat VIP; hiển thị paywall rõ tính năng/quyền lợi.","Có CTA mua VIP và nút quay lại/đóng; không kích hoạt mic trước đồng ý.","Không trừ Tim/EXP."],data="Free",priority="P0")
tc("Tab 2 – AI Voice Chat","VIP gửi và nhận voice message",
   ["Đăng nhập VIP, mở Tab 2.","Bấm mic, thu câu thoại, dừng và gửi.","Chờ phản hồi AI; phát/pause voice message."],
   ["Bubble người dùng/AI đúng phía, thứ tự và timestamp hợp lý.","AI trả voice phù hợp Mode/bối cảnh; audio phát được.","Không trừ Tim; trạng thái VIP vô hạn Tim giữ nguyên."],data="VIP")
tc("Tab 2 – AI Voice Chat","Hủy bản ghi trước khi gửi",
   ["Bắt đầu ghi âm.","Bấm Hủy/Xóa bản nháp.","Quan sát chat."],
   ["Mic dừng; bản ghi không gửi/không tạo bubble.","Không gọi AI, không thay đổi điểm/Tim; có thể ghi lại."],data="VIP",priority="P2")
tc("Tab 2 – AI Voice Chat","Mạng mất khi gửi voice",
   ["Thu voice hợp lệ.","Ngắt mạng trước khi gửi.","Khôi phục mạng và bấm Retry."],
   ["Bubble có trạng thái lỗi/đang chờ; không mất quyền kiểm soát.","Retry gửi đúng một message; AI không trả lời trùng.","Không trừ Tim/cộng EXP vì lỗi truyền."],data="VIP",kind="Khả năng phục hồi")
tc("Tab 2 – AI Voice Chat","VIP hết hạn khi đang chat",
   ["Mở chat khi VIP còn hiệu lực.","Làm VIP hết hạn phía server.","Gửi voice tiếp theo."],
   ["Server từ chối quyền VIP đúng cách; UI hiển thị nâng cấp/gia hạn.","Không gửi tiếp/không thu phí sai; lịch sử đã có xử lý theo chính sách.","Loại gói và Tim cập nhật nhất quán."],data="VIP sắp hết hạn",priority="P0")

tc("Tab 3 – Chính tả","So khớp bỏ qua hoa/thường và dấu câu",
   ["Phát audio một câu.","Nhập đúng chữ nhưng đổi hoa/thường và bỏ/thêm dấu câu.","Bấm “Kiểm tra”."],
   ["Câu được tính đúng theo tỷ lệ nội dung từ; không phạt hoa/thường/dấu câu.","Nếu đạt ngưỡng: +20 EXP một lần và cho câu tiếp.","Hiển thị đáp án/feedback hợp lý."],data="Expected: Hello, world!; Input: hello world")
tc("Tab 3 – Chính tả","Khoảng trắng và ký tự",
   ["Nhập nhiều khoảng trắng đầu/cuối/giữa.","Nhập thiếu từ hoặc thêm từ.","Bấm Kiểm tra từng trường hợp."],
   ["Khoảng trắng được chuẩn hóa; không làm sai câu đúng.","Thiếu/thêm từ ảnh hưởng tỷ lệ theo thuật toán đã công bố; điểm hiển thị nhất quán.","Không lỗi Unicode/apostrophe."],priority="P2")
for threshold in [50,70,90]:
    tc("Tab 3 – Chính tả",f"Ngưỡng đạt chính tả {threshold}%",
       [f"Đặt mục tiêu {threshold}%.",f"Gửi đáp án mock {threshold-1}%.",f"Gửi đáp án {threshold}% ở câu khác."],
       [f"{threshold-1}%: không cho qua, trừ 1 Tim Free.",f"{threshold}%: cho qua và +20 EXP đúng một lần.","Nút Next chỉ enabled sau đạt; VIP không bị trừ khi fail."],
       data=f"Free/VIP; score {threshold-1}/{threshold}")
tc("Tab 3 – Chính tả","Submit rỗng và phát lại audio",
   ["Không nhập gì và bấm Kiểm tra.","Bấm phát lại nhiều lần, đổi tốc độ nếu có.","Nhập đáp án rồi submit."],
   ["Rỗng bị validation và không trừ Tim nếu chưa được coi là attempt.","Phát lại không cộng/trừ điểm; audio đúng câu.","Một submit hợp lệ tạo tối đa một giao dịch."],priority="P2")

tc("Tab 4 – Ghép IPA","Chọn đáp án IPA đúng",
   ["Mở Tab 4.","Đọc từ tiếng Anh và các lựa chọn.","Chọn đáp án đúng; bấm “Kiểm tra/Xác nhận”."],
   ["Chỉ một lựa chọn selected; trạng thái accessible.","Đúng: feedback, +20 EXP đúng một lần, cho câu tiếp.","Không trừ Tim."],data="IPA đúng")
tc("Tab 4 – Ghép IPA","Chọn IPA sai – Free/VIP",
   ["Với Free chọn đáp án sai và xác nhận.","Lặp bằng VIP.","Kiểm tra thanh trạng thái."],
   ["Free trừ đúng 1 Tim, không +EXP; hiển thị đáp án/giải thích theo thiết kế.","VIP không giảm Tim và vẫn không +EXP.","Không cho tính đúng nếu đổi lựa chọn sau khi đã chấm nếu không bấm thử lại."],data="Free và VIP")
tc("Tab 4 – Ghép IPA","Không chọn đáp án và double click",
   ["Không chọn, bấm Xác nhận.","Chọn đúng rồi double click Xác nhận/Next."],
   ["Không chọn: validation, không trừ Tim.","Double click không +40 EXP/không nhảy hai câu.","Nút bị khóa trong lúc xử lý."],priority="P0")
tc("Bài giảng – Chung","Hoàn thành câu cuối/bài học",
   ["Đạt câu cuối ở từng tab áp dụng.","Bấm Hoàn thành/Next.","Mở lại bài và trang Tiến trình."],
   ["Không vượt index/404; hiển thị tổng kết.","EXP/Streak/tiến độ cập nhật chính xác, không cộng lại khi refresh.","Mở lại thể hiện trạng thái hoàn thành theo thiết kế."],priority="P0")

# Modes
tc("Đánh giá theo Mode","Casual chấp nhận dạng rút gọn và slang",
   ["Đặt Mode Giao tiếp.","Dùng gonna/wanna/lemme hoặc slang phù hợp ngữ cảnh.","Gửi đánh giá; so với câu tương đương chuẩn."],
   ["Không phạt máy móc chỉ vì dạng rút gọn/slang hợp lệ.","Feedback ưu tiên fluency và ngữ điệu tự nhiên.","Điểm/giải thích nhất quán, không xúc phạm người học."],data="Bộ audio chuẩn hóa",kind="AI/Chất lượng")
tc("Đánh giá theo Mode","Academic bắt phoneme, trọng âm, ngữ pháp",
   ["Đặt Mode Học thuật.","Gửi mẫu sai phoneme/trọng âm nhưng trôi chảy.","Gửi mẫu dùng ngữ pháp/từ nối tốt."],
   ["Mẫu sai bị chỉ rõ âm/trọng âm/ngữ pháp và chấm nghiêm.","Mẫu tốt được ghi nhận từ vựng nâng cao/từ nối.","Không áp rubric Casual."],data="Bộ audio chuẩn hóa",kind="AI/Chất lượng")
tc("Đánh giá theo Mode","Professional kiểm clarity và sắc thái",
   ["Đặt Mode Công việc.","Nói “I want...” với khách hàng.","Nói lại “I would like...” rõ ràng."],
   ["Mẫu đầu được nhắc về sắc thái/lịch sự; feedback nêu lựa chọn tốt hơn.","Mẫu sau được ghi nhận clarity/tính chuyên nghiệp.","Không đánh đồng hoàn toàn với rubric Academic."],data="Bộ audio chuẩn hóa",kind="AI/Chất lượng")
tc("Đánh giá theo Mode","Cùng một mẫu qua ba Mode",
   ["Giữ cùng audio/text chuẩn hóa.","Chấm độc lập ở Casual, Academic, Professional.","So sánh score breakdown/feedback."],
   ["Rubric thay đổi đúng Mode nhưng kết quả có thể giải thích và tái lập trong dung sai.","Không rò cấu hình Mode cũ; nhãn Mode hiển thị đúng tại thời điểm chấm.","Log lưu rubric/version để truy vết."],kind="AI/Chất lượng")

# Progress
tc("Tiến trình & Thẻ nhớ","Tổng quan Streak/Tim/EXP/gói",
   ["Mở trang Tiến trình.","So sánh với Top Nav và dữ liệu chuẩn.","Refresh/đăng nhập lại."],
   ["Bốn chỉ số hiển thị đúng, có nhãn và format; VIP thể hiện Tim vô hạn.","Đồng nhất trên mọi nơi và bền vững sau phiên mới.","Loading/skeleton không nhấp nháy thành 0 gây hiểu nhầm."],priority="P0")
tc("Tiến trình & Thẻ nhớ","Đổi 100 EXP lấy 1 Tim thành công",
   ["Chuẩn bị Free có đúng 100 EXP.","Bấm “Đổi EXP lấy Tim”.","Đọc popup rồi bấm Xác nhận hai lần nhanh."],
   ["Popup nêu tỷ lệ, số dư trước/sau; chỉ một giao dịch.","EXP giảm 100, Tim tăng 1; tất cả màn hình đồng bộ.","Có thông báo/biên nhận; refresh không đảo/nghịch giao dịch."],data="Free: EXP=100",priority="P0")
tc("Tiến trình & Thẻ nhớ","EXP 99 không đủ đổi Tim",
   ["Chuẩn bị 99 EXP.","Bấm nút đổi hoặc quan sát trạng thái.","Thử gọi hành động qua URL/API."],
   ["Nút disabled hoặc server từ chối với thông báo thiếu 1 EXP.","Không thay đổi EXP/Tim; server không cho âm EXP.","Không thể bypass bằng client."],data="Free: EXP=99",priority="P0",kind="Chức năng/Bảo mật")
tc("Tiến trình & Thẻ nhớ","EXP 101 và hủy popup đổi",
   ["Chuẩn bị 101 EXP.","Bấm đổi rồi bấm Hủy.","Mở lại và xác nhận."],
   ["Hủy giữ 101 EXP/Tim cũ.","Xác nhận còn 1 EXP và +1 Tim.","Focus modal đúng; Escape/đóng xử lý như Hủy."],data="Free: EXP=101")
tc("Tiến trình & Thẻ nhớ","Tỷ lệ đổi EXP cấu hình được",
   ["Đặt tỷ lệ test 120 EXP/Tim phía cấu hình.","Mở trang và giao dịch với 119/120 EXP.","Đổi lại cấu hình 100."],
   ["UI và server cùng dùng 120; 119 bị chặn, 120 thành công.","Không hard-code 100 ở message/validation.","Đổi cấu hình không sửa lịch sử giao dịch cũ."],data="Tỷ lệ 120",kind="Cấu hình")
tc("Tiến trình & Thẻ nhớ","VIP đổi EXP lấy Tim",
   ["Đăng nhập VIP.","Quan sát nút đổi; thử kích hoạt.","Kiểm số dư."],
   ["Vì Tim vô hạn, nút được ẩn/disabled kèm giải thích.","Không trừ EXP, không tạo giao dịch vô nghĩa."],data="VIP",priority="P1")
tc("Tiến trình & Thẻ nhớ","Xóa câu thoại yêu thích",
   ["Mở danh sách câu yêu thích.","Bấm Xóa; bấm Hủy.","Xóa lại và Xác nhận; refresh."],
   ["Hủy giữ câu; xác nhận chỉ xóa câu chọn.","Danh sách/empty state cập nhật; không xóa flashcard liên quan ngoài PRD.","Sau refresh câu vẫn đã xóa."],priority="P1")
tc("Flashcard","Hiển thị nội dung flashcard",
   ["Mở Sổ tay từ vựng.","Duyệt card được tự thêm và card khác.","Lật card nếu có."],
   ["Hiển thị từ, IPA, nghĩa và nguồn/ngữ cảnh cần thiết; không lộ đáp án sai.","Text dài wrap, không cắt; card thao tác bằng bàn phím/touch.","Không trùng card cho cùng từ theo quy tắc sản phẩm."],priority="P1")
tc("Flashcard","Hành động Đã nhớ",
   ["Tại card bấm “Đã nhớ”.","Kiểm tra trạng thái/card kế tiếp.","Refresh."],
   ["Card được đánh dấu đã nhớ/di chuyển theo thiết kế; feedback rõ.","Không bị xóa vĩnh viễn nếu PRD chỉ yêu cầu trạng thái; trạng thái bền vững.","Không đổi EXP/Tim nếu không quy định."],priority="P1")
tc("Flashcard","Bật mic thu âm đọc lại",
   ["Tại card bấm icon Mic.","Cho phép mic, đọc từ và dừng/gửi.","Thử Block mic."],
   ["Thu/chấm đúng từ; hiển thị feedback phát âm nếu có.","Mic dừng khi xong/rời card; Block có hướng dẫn và không trừ Tim ngoài PRD.","Không tự đánh dấu Đã nhớ nếu chưa có quy tắc."],priority="P1")
tc("Flashcard","Hành động Học lại",
   ["Bấm “Học lại” trên flashcard.","Thực hiện bài luyện mở ra.","Dùng Back."],
   ["Mở đúng bài/câu/từ liên quan hoặc bài luyện phù hợp.","Không mất flashcard; Back quay đúng vị trí.","Điểm/Tim chỉ thay đổi theo hoạt động chấm thực sự."],priority="P1")

# Account
tc("Tài khoản","Hiển thị thông tin cá nhân",
   ["Mở Tài khoản.","Đối chiếu Tên, Số điện thoại, Email."],
   ["Đúng người dùng; không lộ thông tin người khác khi đổi tài khoản.","Email hiển thị readonly/disabled và không thể sửa bằng bàn phím.","Dữ liệu nhạy cảm không xuất hiện trong URL."],priority="P0")
tc("Tài khoản","Sửa Tên và Số điện thoại hợp lệ",
   ["Bấm Sửa.","Đổi Tên và Số điện thoại.","Bấm “Lưu”; refresh và đăng nhập lại."],
   ["Validation hợp lệ; loading/chống double click.","Thông báo thành công; dữ liệu mới hiển thị và bền vững.","Không thay đổi Email/Mode/điểm."],data="Tên Unicode; SĐT hợp lệ")
tc("Tài khoản","Validation Tên/Số điện thoại",
   ["Nhập tên trống/chỉ khoảng trắng/quá dài.","Nhập SĐT chữ, quá ngắn/dài, ký tự không hợp lệ.","Bấm Lưu."],
   ["Thông báo sát trường; không gửi hoặc server từ chối nhất quán.","Không lưu một phần; dữ liệu cũ vẫn nguyên.","Không thực thi HTML/script trong tên."],priority="P0",kind="Validation/Bảo mật")
tc("Tài khoản","Không thể đổi Email bằng client/API",
   ["Thử focus/sửa trường Email.","Sửa DOM/request để gửi email khác cùng form.","Refresh."],
   ["UI không cho sửa; server bỏ qua/từ chối trường email trái phép.","Email gốc không đổi; có log sự kiện nghi ngờ phù hợp.","Không hiển thị stack trace."],priority="P0",kind="Bảo mật")
tc("Cấu hình học tập","Đổi ngưỡng 50/70/90 và lưu",
   ["Lần lượt chọn 50%, 70%, 90%; mỗi lần bấm Lưu.","Refresh sau mỗi lần.","Mở Tab 1/3 và kiểm ngưỡng."],
   ["Chỉ một giá trị selected; Lưu thành công và bền vững.","Tab 1/3 dùng ngưỡng mới cho lần chấm sau; không sửa lịch sử cũ.","Không tự đổi Mode/giọng Anh."],priority="P0")
tc("Cấu hình học tập","Đổi chuẩn Anh-Anh / Anh-Mỹ",
   ["Chọn Anh-Anh và Lưu; mở bài có audio/IPA.","Chọn Anh-Mỹ và Lưu; mở lại.","Refresh/đăng nhập lại."],
   ["Chỉ một chuẩn active; nội dung/IPA/audio/đánh giá dùng chuẩn đã chọn khi dữ liệu hỗ trợ.","Không trộn chuẩn trong cùng lần chấm; cấu hình bền vững.","Nếu bài không hỗ trợ, có thông báo/fallback rõ."],priority="P1")
tc("Đổi Mode","Free đổi Mode – popup xác nhận",
   ["Đăng nhập Free còn lượt đổi.","Chọn Mode khác.","Trong popup đọc ảnh hưởng rồi bấm Hủy.","Lặp lại và bấm Xác nhận."],
   ["Hủy giữ Mode cũ và không mất lượt/EXP.","Xác nhận đổi đúng Mode, cập nhật khóa học/rubric và số lượt/EXP theo cấu hình.","Popup chặn click nền, focus trap, Escape=Hủy."],data="Free còn lượt",priority="P0")
tc("Đổi Mode","Free hết lượt đổi theo giới hạn tháng",
   ["Chuẩn bị Free đã dùng lượt tháng.","Chọn Mode khác và xác nhận.","Thử bypass qua request."],
   ["UI chặn và nêu ngày/lượt có lại hoặc lựa chọn mua lượt.","Server không đổi Mode; không trừ EXP ngoài ý muốn.","Mode và rubric vẫn cũ trên toàn app."],data="Free hết lượt",priority="P0")
tc("Đổi Mode","Free mua lượt đổi bằng EXP",
   ["Bật cấu hình đổi Mode bằng EXP.","Chuẩn bị đủ EXP; chọn Mode và xác nhận chi phí.","Double click Xác nhận."],
   ["Trừ đúng chi phí một lần, đổi Mode một lần; hiển thị trước/sau.","Không đủ EXP thì chặn server-side và không đổi.","Giao dịch bền vững, không thể replay."],data="[CẦN PO XÁC NHẬN] chi phí",priority="P0")
tc("Đổi Mode","VIP đổi Mode không giới hạn",
   ["Đăng nhập VIP.","Đổi Giao tiếp → Học thuật → Công việc liên tiếp, xác nhận mỗi lần.","Refresh."],
   ["Không giới hạn/không trừ EXP; Mode cuối là Công việc.","Khóa học, AI prompt role và rubric cập nhật theo Mode mới.","Không tạo lỗi race khi đổi nhanh."],data="VIP",priority="P0")
tc("Tài khoản Mobile","Đổi Dark/Light trong Tài khoản",
   ["Đặt viewport Mobile, mở Tài khoản.","Bấm công tắc theme.","Đi qua các trang và refresh."],
   ["Theme đổi; công tắc không xuất hiện trùng ở thanh trên.","Trạng thái lưu đúng; icon/text/tương phản rõ.","Không gây scroll ngang/nhảy layout lớn."],kind="Responsive")
tc("Tài khoản Mobile","Đăng xuất trong Tài khoản",
   ["Đặt viewport Mobile, mở Tài khoản.","Cuộn tới Đăng xuất và bấm.","Xác nhận; dùng Back."],
   ["Đăng xuất thành công; không còn dữ liệu cá nhân.","Không có nút Đăng xuất trùng ở thanh điều hướng Mobile.","Back không mở trang bảo vệ."],priority="P0",kind="Responsive/Bảo mật")
tc("VIP","Mở luồng Mua VIP",
   ["Free mở Tài khoản.","Bấm “Mua VIP”.","Chọn MoMo hoặc ZaloPay và tiếp tục."],
   ["Hiển thị gói/giá/quyền lợi/chu kỳ/gia hạn minh bạch.","Tạo giao dịch đúng tài khoản/số tiền; chuyển cổng đúng.","Chưa có callback thành công thì vẫn Free."],priority="P0")
tc("VIP","Thanh toán VIP thành công",
   ["Khởi tạo thanh toán.","Giả lập callback thành công hợp lệ.","Quay lại app/refresh."],
   ["Kích hoạt VIP đúng một lần; Top Nav/Tiến trình cập nhật gói và Tim vô hạn.","Mở được Tab 2; quyền đổi Mode không giới hạn.","Callback lặp không gia hạn/kích hoạt trùng; có biên nhận phù hợp."],priority="P0",kind="Tích hợp thanh toán")
tc("VIP","Thanh toán thất bại/hủy/timeout",
   ["Khởi tạo thanh toán.","Hủy tại cổng; lặp với thất bại và pending timeout.","Quay lại app."],
   ["Trạng thái rõ ràng; không kích hoạt VIP khi chưa thành công.","Có thể thử lại an toàn; không tạo nhiều khoản thu do retry.","Không mất cấu hình/tiến độ."],priority="P0",kind="Tích hợp thanh toán")
tc("VIP","Bảo mật callback thanh toán",
   ["Gửi callback chữ ký sai/số tiền sai/tài khoản sai.","Gửi replay callback hợp lệ cũ.","Kiểm tra trạng thái."],
   ["Server từ chối callback không hợp lệ, log/audit an toàn.","Replay idempotent; không cộng thời hạn/quyền lợi trùng.","Không lộ secret/chữ ký/token trong UI/log client."],priority="P0",kind="Bảo mật")
tc("VIP","Hủy gia hạn VIP",
   ["VIP bấm “Hủy gia hạn”.","Bấm Hủy trong popup; lặp lại và Xác nhận.","Refresh/đăng nhập lại."],
   ["Lần đầu không thay đổi; xác nhận tắt gia hạn nhưng giữ VIP tới ngày hết hạn.","Hiển thị ngày kết thúc/trạng thái rõ; có thể mua/gia hạn lại theo thiết kế.","Không xóa tiến độ hoặc trừ Tim hồi tố."],priority="P0")

# Responsive
for size, label in [("1920×1080","Desktop lớn"),("1366×768","Laptop"),("1024×768","Desktop nhỏ/Tablet ngang"),("768×1024","Tablet dọc"),("390×844","Mobile chuẩn"),("320×568","Mobile nhỏ")]:
    tc("Responsive",f"Bố cục toàn app tại {label} {size}",
       [f"Đặt viewport {size}, zoom 100%.","Đi qua Đăng nhập, Khóa học, 4 Tab bài giảng, Tiến trình, Tài khoản.","Mở dropdown/modal/keyboard ảo nếu là mobile."],
       ["Không scroll ngang ngoài vùng nội dung có chủ đích; không chồng/cắt chữ/nút.","Điều hướng đúng biến thể Desktop/Mobile theo breakpoint; nội dung chính không bị che.","Nút có vùng bấm đủ lớn; modal nằm trong viewport và cuộn nội dung được; media giữ tỷ lệ."],
       priority="P1",kind="Responsive")
tc("Responsive","Chuyển qua lại breakpoint không mất trạng thái",
   ["Đang nhập prompt/đáp án ở 1024 px.","Resize 1024 → 767 → 768 → 390 → 1440.","Tiếp tục thao tác."],
   ["Nav chuyển đúng tại biên, không hiện đồng thời hai loại.","Không mất input/bản ghi đã dừng/trạng thái tab; không request lặp.","Không để overlay/backdrop kẹt."],kind="Responsive")
tc("Responsive","Xoay màn hình và bàn phím ảo",
   ["Trên mobile mở field prompt/chính tả để bật bàn phím.","Xoay portrait ↔ landscape.","Đóng bàn phím và submit."],
   ["Field/nút submit không bị che; tự cuộn hợp lý.","Không reload/mất dữ liệu; viewport và modal tính lại đúng.","Submit một lần và focus không nhảy vô hạn."],kind="Responsive")
tc("Responsive","Zoom 200% và reflow",
   ["Desktop đặt zoom 200%.","Đi qua toàn app và mở modal.","Dùng chỉ bàn phím."],
   ["Nội dung reflow, không mất chức năng/không cần scroll hai chiều bất hợp lý.","Text không cắt; focus luôn nhìn thấy; modal thao tác được.","Thông tin không chỉ truyền bằng vị trí/màu."],kind="Accessibility/Responsive")

# Non-functional
nonfunc = [
("Hiệu năng tải trang","Đo Cold/Warm load Trang Khóa Học với dữ liệu đại diện.",["Có skeleton phù hợp; LCP/INP/CLS đạt ngân sách dự án [CẦN PO CHỐT].","Không tải toàn bộ audio/video ngoài viewport; tương tác đầu không bị khóa."]),
("Hiệu năng danh sách lớn","Nạp ≥500 bài, 1.000 câu yêu thích/flashcard và cuộn/lọc.",["Phân trang/virtualization hợp lý; cuộn mượt, không tăng RAM vô hạn.","Không trùng/mất item; thứ tự ổn định."]),
("Tải đồng thời","Mô phỏng người dùng đồng thời tạo AI/chấm phát âm/thanh toán theo tải mục tiêu.",["Tỷ lệ lỗi/latency trong SLO [CẦN PO CHỐT].","Không sai EXP/Tim/quyền VIP do race."]),
("Khả năng phục hồi mạng","Mất mạng giữa lưu cấu hình, chấm điểm và giao dịch đổi EXP.",["Không cập nhật optimistic sai kéo dài; retry/idempotency đúng.","Thông báo trạng thái và khôi phục không tạo giao dịch kép."]),
("Bảo mật xác thực","Thử truy cập IDOR trang/bài/tiến trình tài khoản khác; token hết hạn/giả mạo.",["401/403 đúng; không lộ PII/nội dung riêng.","Token ở cookie/header an toàn; không nằm URL; logout vô hiệu hóa phiên."]),
("XSS/Injection","Đưa payload vào Tên, prompt, transcript, AI output và query.",["Encode/sanitize mọi nơi; không thực thi script/HTML nguy hiểm.","API/DB không bị injection; lỗi không lộ stack/query."]),
("CSRF/CORS","Gửi request đổi cấu hình, xóa, thanh toán từ origin lạ.",["Chặn theo cơ chế CSRF/CORS phù hợp; chỉ origin tin cậy.","Không thay đổi trạng thái khi thiếu token/chứng thực."]),
("Quyền riêng tư audio","Kiểm network/log/storage khi ghi và gửi audio.",["Chỉ thu khi có chỉ báo/đồng ý; dừng đúng lúc.","Không lưu audio cục bộ/server quá thời hạn chính sách; truyền TLS; có thông báo chính sách."]),
("Accessibility bàn phím","Dùng Tab/Shift+Tab/Enter/Space/Escape cho mọi màn.",["Thứ tự focus logic; focus visible; không keyboard trap.","Modal trap focus và trả focus; media/tab/radio hoạt động."]),
("Accessibility screen reader","Đọc Logo, nav, status, form error, tab, score, flashcard, modal.",["Tên/role/state/aria-live đúng; lỗi liên kết trường.","Icon-only có accessible name; thay đổi điểm được thông báo không gây spam."]),
("Tương phản & màu","Kiểm Light/Dark, hover/focus/disabled/error/success.",["Tương phản đạt WCAG AA mục tiêu; trạng thái không chỉ dựa màu.","Dark mode không còn vùng sáng/chữ chìm."]),
("Đa ngôn ngữ/ký tự","Dùng tên Việt có dấu, IPA, apostrophe, emoji, chuỗi dài.",["Lưu/hiển thị/tìm kiếm không lỗi mã hóa; IPA đúng glyph.","Wrap/ellipsis có tooltip hợp lý; không vỡ DB/export."]),
("Tính nhất quán dữ liệu","Mở cùng tài khoản ở hai tab/thiết bị, đồng thời học/đổi EXP/Mode.",["Không lost update/âm điểm; server là nguồn chuẩn.","UI phát hiện/cập nhật thay đổi; giao dịch có tính nguyên tử."]),
("Quan sát & audit","Gây lỗi AI, pronunciation API, payment, đổi Mode, đổi EXP.",["Log có correlation ID, user pseudonymous, latency/error code; không chứa mật khẩu/token/audio nhạy cảm.","Audit đủ truy vết thay đổi điểm/gói/cấu hình."]),
("Khả năng tương thích","Chạy smoke trên Chrome/Edge/Firefox/Safari mục tiêu.",["Tính năng chính và media/mic hoạt động nhất quán; fallback rõ.","Không dùng API trình duyệt không hỗ trợ mà thiếu kiểm tra."]),
("Khôi phục sau refresh/crash","Refresh/đóng tab giữa recording, payment pending, AI generation.",["Không gửi audio/bài/giao dịch nửa vời; trạng thái pending được hòa giải.","Dữ liệu đã commit không mất; thao tác chưa commit không tự tính điểm."]),
]
for title, action, exp in nonfunc:
    tc("Phi chức năng",title,[action,"Ghi nhận metric/log và lặp tối thiểu 3 lần khi phù hợp."],exp,priority="P0" if any(k in title for k in ["Bảo mật","Quyền riêng tư","Tính nhất quán"]) else "P1",kind="Phi chức năng")

doc.add_page_break()
doc.add_heading("5. Danh sách kịch bản kiểm thử chi tiết", level=1)
p=doc.add_paragraph()
set_font(p.add_run(f"Tổng số test case: {len(cases)}. "),11,True,NAVY)
set_font(p.add_run("Thứ tự ưu tiên: P0 = chặn phát hành; P1 = chức năng cốt lõi; P2 = quan trọng nhưng có workaround."),11)

areas=[]
for c in cases:
    if c["area"] not in areas: areas.append(c["area"])
summary=doc.add_table(rows=1,cols=3); summary.style="Table Grid"; set_table_widths(summary,[0.7,4.8,1.0])
for i,h in enumerate(["STT","Phân hệ","Số TC"]):
    summary.cell(0,i).text=h; set_cell_shading(summary.cell(0,i),LIGHT)
    for r in summary.cell(0,i).paragraphs[0].runs:set_font(r,10,True,NAVY)
set_repeat_table_header(summary.rows[0])
for idx,a in enumerate(areas,1):
    n=sum(1 for c in cases if c["area"]==a)
    row=summary.add_row().cells; row[0].text=str(idx); row[1].text=a; row[2].text=str(n)
    for cell in row:
        for r in cell.paragraphs[0].runs:set_font(r,9.5)
    set_table_widths(summary,[0.7,4.8,1.0])

current=None
area_num=0
for idx,c in enumerate(cases,1):
    if c["area"] != current:
        current=c["area"]; area_num += 1
        doc.add_page_break()
        doc.add_heading(f"5.{area_num}. {current}",level=2)
    case_id=f"TC-{idx:03d}"
    h=doc.add_paragraph(style="Heading 3")
    h.paragraph_format.page_break_before = False
    set_font(h.add_run(f"{case_id}  |  {c['title']}"),12,True,NAVY)
    m=doc.add_table(rows=3,cols=4); m.style="Table Grid"; set_table_widths(m,[0.9,2.35,0.9,2.35])
    meta_rows=[
        ("Ưu tiên",c["priority"],"Loại",c["kind"]),
        ("Tiền điều kiện",c["pre"],"Dữ liệu",c["data"]),
        ("Kết quả thực tế","(Tester ghi khi thực thi)","Trạng thái","☐ Pass  ☐ Fail  ☐ Blocked  ☐ N/A"),
    ]
    for ri,rowv in enumerate(meta_rows):
        for ci,val in enumerate(rowv):
            m.cell(ri,ci).text=val
            if ci in (0,2): set_cell_shading(m.cell(ri,ci),LIGHT)
            for r in m.cell(ri,ci).paragraphs[0].runs: set_font(r,9,ci in (0,2),NAVY if ci in (0,2) else None)
    pl=doc.add_paragraph(style="TC Label"); pl.add_run("CÁC BƯỚC THỰC HIỆN")
    for s in c["steps"]:
        doc.add_paragraph(s,style="List Number")
    pl=doc.add_paragraph(style="TC Label"); pl.add_run("KẾT QUẢ MONG ĐỢI")
    for e in c["expected"]:
        doc.add_paragraph(e,style="List Bullet")
    sep=doc.add_paragraph()
    sep.paragraph_format.space_after=Pt(2)
    pPr=sep._p.get_or_add_pPr()
    pBdr=OxmlElement("w:pBdr"); bottom=OxmlElement("w:bottom")
    bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"4"); bottom.set(qn("w:space"),"1"); bottom.set(qn("w:color"),"D9E2F3")
    pBdr.append(bottom); pPr.append(pBdr)

doc.add_page_break()
doc.add_heading("6. Checklist hồi quy phát hành",level=1)
for text in [
    "Đăng ký đủ 3 Mode × 3 ngưỡng; đăng nhập/đăng xuất/phiên hết hạn.",
    "Desktop Top Nav và Mobile Sub-nav đúng; Dark/Light bền vững.",
    "Giáo trình và Video Bank đủ 3 phân loại; card/empty/error state.",
    "AI lesson: tạo, không lưu, lưu, auto-save, xóa, lỗi/timeout và nội dung theo Mode.",
    "4 Tabs; Free paywall Tab 2; mic allow/block; audio/transcript/IPA/nghĩa.",
    "Boundary 49/50/51, 69/70/71, 89/90/91; +20 EXP; -1 Tim; VIP vô hạn.",
    "Sai từ lần thứ 4 tự vào Sổ tay; lưu/xóa câu yêu thích; 3 hành động flashcard.",
    "Đổi 100 EXP lấy Tim, thiếu EXP, double click, tỷ lệ cấu hình.",
    "Sửa Tên/SĐT, Email readonly/server protected, Anh-Anh/Anh-Mỹ, đổi Mode Free/VIP.",
    "Mua/Hủy gia hạn VIP; payment success/fail/pending/replay/signature invalid.",
    "Responsive 1920→320, breakpoint 767/768/1023/1024, zoom 200%, keyboard/screen reader.",
    "Security, privacy audio, performance, concurrency, idempotency, logging và recovery.",
]:
    doc.add_paragraph("☐ " + text,style="List Bullet")

doc.add_heading("7. Mẫu ghi nhận lỗi",level=1)
for label, value in [
    ("Tiêu đề","[Phân hệ] Mô tả ngắn sai lệch"),
    ("Môi trường","Thiết bị / OS / trình duyệt / viewport / tài khoản / Mode / ngưỡng / loại gói"),
    ("Tiền điều kiện","Số dư EXP/Tim, dữ liệu bài, quyền mic, cấu hình feature flag"),
    ("Bước tái hiện","Danh sách thao tác tối thiểu, có giá trị nhập và thời điểm"),
    ("Thực tế vs mong đợi","Kèm số liệu trước/sau; nêu bất biến bị phá"),
    ("Bằng chứng","Ảnh/video, HAR, console/server log, correlation ID; che dữ liệu nhạy cảm"),
    ("Mức độ","Blocker/Critical/Major/Minor và ảnh hưởng người dùng"),
]:
    p=doc.add_paragraph()
    set_font(p.add_run(label + ": "),11,True,NAVY)
    set_font(p.add_run(value),11)

doc.core_properties.title = "Kịch bản kiểm thử chi tiết PRD WebShadowing"
doc.core_properties.subject = "User flow, responsive, functional and non-functional test scenarios"
doc.core_properties.author = "Codex"
doc.core_properties.keywords = "WebShadowing, PRD, test case, responsive, VIP, pronunciation"
doc.save(OUT)
print(f"CREATED={OUT.resolve()}")
print(f"TEST_CASES={len(cases)}")
